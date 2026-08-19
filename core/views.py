from pathlib import Path
import json
import re

import requests

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.shortcuts import redirect, render

from pypdf import PdfReader
from docx import Document

from .models import CareerSearch, ResumeAnalysis


# ============================================================
# OLLAMA / LOCAL GEMMA AI
# ============================================================

OLLAMA_URL = "http://127.0.0.1:11434/api/generate"
OLLAMA_MODEL = "gemma3:4b"


def ask_gemma(prompt):
    print("\n====================================")
    print("GEMMA REQUEST STARTED")
    print("====================================")

    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
            },
            timeout=180,
        )

        response.raise_for_status()

        data = response.json()
        result = data.get("response", "").strip()

        print("GEMMA RESPONSE RECEIVED")
        print("Response length:", len(result))
        print("====================================\n")

        return result

    except requests.exceptions.ConnectionError:
        print("OLLAMA ERROR: Ollama is not running.")
        return ""

    except requests.exceptions.Timeout:
        print("OLLAMA ERROR: Request timed out.")
        return ""

    except Exception as e:
        print("OLLAMA ERROR:", e)
        return ""


# ============================================================
# HOME / CAREER GUIDANCE
# ============================================================

def home(request):

    if request.method == "POST":

        career = request.POST.get("career", "").strip()

        if not career:
            return render(
                request,
                "home.html",
                {"error": "Please enter a career."}
            )

        print("\n====================================")
        print("CAREER SEARCH:", career)
        print("====================================")
        print("Sending career guidance request to Gemma...")

        career_info = get_career_information(career)

        # ----------------------------------------------------
        # SAVE SEARCH
        # ----------------------------------------------------

        if request.user.is_authenticated:

            try:
                CareerSearch.objects.create(
                    user=request.user,
                    career=career,
                    source="Career Assistant"
                )

            except Exception as e:
                print("CAREER SEARCH SAVE ERROR:", e)

        # ----------------------------------------------------
        # GEMMA CAREER ANALYSIS
        # ----------------------------------------------------

        ai_analysis = generate_career_ai_analysis(career)

        return render(
            request,
            "home.html",
            {
                "career": career,
                "career_info": career_info,
                "ai_analysis": ai_analysis,
                "error": None,
            }
        )

    return render(request, "home.html")


# ============================================================
# CAREER DATABASE
# ============================================================

CAREER_DATABASE = {

    "python developer": {
        "skills": [
            "Python",
            "Django",
            "SQL",
            "Git",
            "REST API"
        ],
        "projects": [
            "Django Job Portal",
            "REST API Project",
            "Task Management System"
        ],
        "roadmap": [
            "Master Python",
            "Learn Django",
            "Learn SQL",
            "Learn Git and GitHub",
            "Build REST APIs",
            "Build full-stack projects"
        ],
        "interview": [
            "Python fundamentals",
            "Object-Oriented Programming",
            "Django",
            "SQL",
            "REST APIs",
            "Git"
        ]
    },

    "full stack developer": {
        "skills": [
            "HTML",
            "CSS",
            "JavaScript",
            "React",
            "Node.js",
            "SQL",
            "Git"
        ],
        "projects": [
            "E-commerce Website",
            "Social Media Application",
            "Job Portal"
        ],
        "roadmap": [
            "Learn HTML and CSS",
            "Learn JavaScript",
            "Learn a frontend framework",
            "Learn backend development",
            "Learn databases",
            "Learn REST APIs",
            "Deploy applications"
        ],
        "interview": [
            "HTML",
            "CSS",
            "JavaScript",
            "React",
            "Backend development",
            "Databases"
        ]
    },

    "data analyst": {
        "skills": [
            "Python",
            "SQL",
            "Excel",
            "Power BI",
            "Pandas",
            "NumPy"
        ],
        "projects": [
            "Sales Dashboard",
            "Customer Analysis",
            "Business Intelligence Dashboard"
        ],
        "roadmap": [
            "Learn Excel",
            "Learn SQL",
            "Learn Python",
            "Learn Pandas",
            "Learn Data Visualization",
            "Learn Power BI",
            "Build portfolio projects"
        ],
        "interview": [
            "SQL",
            "Excel",
            "Python",
            "Statistics",
            "Power BI",
            "Data Visualization"
        ]
    },

    "data scientist": {
        "skills": [
            "Python",
            "SQL",
            "Pandas",
            "NumPy",
            "Machine Learning"
        ],
        "projects": [
            "House Price Prediction",
            "Customer Churn Prediction",
            "Recommendation System"
        ],
        "roadmap": [
            "Learn Python",
            "Learn Statistics",
            "Learn SQL",
            "Learn Pandas and NumPy",
            "Learn Machine Learning",
            "Learn Deep Learning",
            "Build portfolio projects"
        ],
        "interview": [
            "Python",
            "Statistics",
            "Machine Learning",
            "SQL",
            "Pandas",
            "Model evaluation"
        ]
    },

    "machine learning engineer": {
        "skills": [
            "Python",
            "Machine Learning",
            "Deep Learning",
            "TensorFlow",
            "SQL"
        ],
        "projects": [
            "Image Classification",
            "Recommendation System",
            "Prediction Model"
        ],
        "roadmap": [
            "Learn Python",
            "Learn Mathematics",
            "Learn Machine Learning",
            "Learn Deep Learning",
            "Learn TensorFlow or PyTorch",
            "Learn Model Deployment"
        ],
        "interview": [
            "Machine Learning",
            "Deep Learning",
            "Python",
            "Statistics",
            "Model deployment"
        ]
    },

    "cybersecurity": {
        "skills": [
            "Cybersecurity",
            "Linux",
            "Networking",
            "Python"
        ],
        "projects": [
            "Network Security Scanner",
            "Password Security Checker",
            "Log Analysis Tool"
        ],
        "roadmap": [
            "Learn Networking",
            "Learn Linux",
            "Learn Python",
            "Learn Cybersecurity Fundamentals",
            "Learn Web Security",
            "Learn Security Tools"
        ],
        "interview": [
            "Networking",
            "Linux",
            "Cybersecurity",
            "OWASP",
            "Security Fundamentals"
        ]
    },

    "cybersecurity analyst": {
        "skills": [
            "Cybersecurity",
            "Linux",
            "Networking",
            "Python"
        ],
        "projects": [
            "Security Log Analyzer",
            "Network Monitoring Tool",
            "Vulnerability Scanner"
        ],
        "roadmap": [
            "Learn Networking",
            "Learn Linux",
            "Learn Python",
            "Learn Cybersecurity Fundamentals",
            "Learn SIEM concepts",
            "Learn Incident Response"
        ],
        "interview": [
            "Networking",
            "Linux",
            "Cybersecurity",
            "Incident Response",
            "Security Monitoring"
        ]
    },

    "cloud engineer": {
        "skills": [
            "AWS",
            "Linux",
            "Networking",
            "Docker",
            "Git"
        ],
        "projects": [
            "Deploy Django Application",
            "Dockerized Web Application",
            "Cloud Monitoring Project"
        ],
        "roadmap": [
            "Learn Linux",
            "Learn Networking",
            "Learn Cloud Fundamentals",
            "Learn AWS or Azure",
            "Learn Docker",
            "Learn CI/CD"
        ],
        "interview": [
            "Cloud Computing",
            "AWS",
            "Linux",
            "Networking",
            "Docker"
        ]
    },

    "devops engineer": {
        "skills": [
            "Linux",
            "Git",
            "Docker",
            "AWS",
            "Python"
        ],
        "projects": [
            "CI/CD Pipeline",
            "Docker Deployment",
            "Cloud Infrastructure Project"
        ],
        "roadmap": [
            "Learn Linux",
            "Learn Git",
            "Learn Docker",
            "Learn CI/CD",
            "Learn Cloud",
            "Learn Infrastructure Automation"
        ],
        "interview": [
            "Linux",
            "Git",
            "Docker",
            "CI/CD",
            "Cloud"
        ]
    },

    "chartered accountant": {
        "skills": [
            "Financial Accounting",
            "Auditing",
            "Taxation",
            "Financial Analysis",
            "Excel"
        ],
        "projects": [
            "Financial Statement Analysis",
            "Cash Flow Forecasting",
            "Financial Transaction Management System"
        ],
        "roadmap": [
            "Learn accounting fundamentals",
            "Learn financial reporting",
            "Learn auditing",
            "Learn taxation",
            "Gain practical experience",
            "Prepare for the relevant professional qualification"
        ],
        "interview": [
            "Accounting Principles",
            "Financial Statement Analysis",
            "Auditing",
            "Taxation",
            "Financial Reporting"
        ]
    }
}


# ============================================================
# GET CAREER INFORMATION
# ============================================================

def get_career_information(career):

    normalized = career.lower().strip()

    if normalized in CAREER_DATABASE:
        return CAREER_DATABASE[normalized]

    for name, data in CAREER_DATABASE.items():

        if name in normalized or normalized in name:
            return data

    return {
        "skills": [
            "Core fundamentals",
            "Problem Solving",
            "Communication",
            "Technical Skills",
            "Domain Knowledge"
        ],
        "projects": [
            f"{career} Portfolio Project",
            f"{career} Practical Application",
            f"{career} Capstone Project"
        ],
        "roadmap": [
            f"Learn {career} fundamentals",
            "Develop relevant technical skills",
            "Build practical projects",
            "Create a portfolio",
            "Prepare for interviews",
            "Apply for internships and jobs"
        ],
        "interview": [
            f"{career} fundamentals",
            "Problem solving",
            "Projects",
            "Technical concepts",
            "Behavioral questions"
        ]
    }


# ============================================================
# CLEAN GEMMA JSON
# ============================================================

def clean_gemma_json(response):

    if not response:
        return ""

    cleaned = response.strip()

    cleaned = re.sub(
        r"```json",
        "",
        cleaned,
        flags=re.IGNORECASE
    )

    cleaned = re.sub(
        r"```",
        "",
        cleaned
    )

    cleaned = cleaned.strip()

    start = cleaned.find("{")
    end = cleaned.rfind("}")

    if start != -1 and end != -1 and end > start:
        cleaned = cleaned[start:end + 1]

    return cleaned.strip()


# ============================================================
# GEMMA CAREER ANALYSIS
# ============================================================

def generate_career_ai_analysis(career):

    career_info = get_career_information(career)

    prompt = f"""
You are CareerAI, an AI career guidance assistant.

The user searched for this career:

{career}

Your task is to provide useful and realistic career guidance.

Known information for this career:

Skills:
{", ".join(career_info["skills"])}

Projects:
{", ".join(career_info["projects"])}

Roadmap:
{", ".join(career_info["roadmap"])}

Interview topics:
{", ".join(career_info["interview"])}

IMPORTANT:

- The career may be outside the provided database.
- Do not assume the user is a Computer Science student unless relevant.
- Do not invent facts about the user.
- Explain the career clearly.
- Give practical skills.
- Give a realistic learning roadmap.
- Give useful beginner projects.
- Give interview preparation topics.
- If the career requires a professional qualification, mention that.
- Keep the response useful for a student or beginner.

Return ONLY valid JSON.

Use exactly this structure:

{{
    "career": "{career}",
    "overview": "Short explanation of the career.",
    "why_choose": "Why someone may choose this career.",
    "skills": [
        "skill 1",
        "skill 2",
        "skill 3",
        "skill 4",
        "skill 5"
    ],
    "roadmap": [
        "step 1",
        "step 2",
        "step 3",
        "step 4",
        "step 5"
    ],
    "projects": [
        "project 1",
        "project 2",
        "project 3"
    ],
    "interview_topics": [
        "topic 1",
        "topic 2",
        "topic 3",
        "topic 4",
        "topic 5"
    ]
}}

Do not include markdown.
Do not include code fences.
"""

    response = ask_gemma(prompt)

    if not response:

        return {
            "career": career,
            "overview": (
                f"{career} is a professional career that "
                "requires relevant technical knowledge, "
                "problem-solving ability and practical "
                "experience."
            ),
            "why_choose": (
                "This career can provide opportunities "
                "to develop valuable professional skills "
                "and work on real-world problems."
            ),
            "skills": career_info["skills"],
            "roadmap": career_info["roadmap"],
            "projects": career_info["projects"],
            "interview_topics": career_info["interview"],
            "ai_available": False
        }

    try:

        cleaned = clean_gemma_json(response)
        data = json.loads(cleaned)

        data.setdefault("career", career)
        data.setdefault("overview", career_info["roadmap"][0])
        data.setdefault(
            "why_choose",
            "This career offers opportunities for professional growth."
        )
        data.setdefault("skills", career_info["skills"])
        data.setdefault("roadmap", career_info["roadmap"])
        data.setdefault("projects", career_info["projects"])
        data.setdefault(
            "interview_topics",
            career_info["interview"]
        )

        data["ai_available"] = True

        return data

    except Exception as e:

        print("GEMMA JSON ERROR:", e)
        print("GEMMA RAW RESPONSE:")
        print(response)

        return {
            "career": career,
            "overview": (
                f"{career} is a career that requires "
                "relevant professional knowledge and "
                "practical experience."
            ),
            "why_choose": (
                "This career can provide opportunities "
                "to develop useful professional skills."
            ),
            "skills": career_info["skills"],
            "roadmap": career_info["roadmap"],
            "projects": career_info["projects"],
            "interview_topics": career_info["interview"],
            "ai_available": False
        }


# ============================================================
# REGISTER
# ============================================================

def register_view(request):

    if request.method == "POST":

        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip()
        password = request.POST.get("password", "")
        confirm_password = request.POST.get(
            "confirm_password",
            ""
        )

        if not username or not password:

            return render(
                request,
                "register.html",
                {
                    "error":
                        "Username and password are required."
                }
            )

        if password != confirm_password:

            return render(
                request,
                "register.html",
                {
                    "error":
                        "Passwords do not match."
                }
            )

        if User.objects.filter(username=username).exists():

            return render(
                request,
                "register.html",
                {
                    "error":
                        "Username already exists."
                }
            )

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password
        )

        login(request, user)

        return redirect("home")

    return render(request, "register.html")


# ============================================================
# LOGIN
# ============================================================

def login_view(request):

    if request.method == "POST":

        username = request.POST.get(
            "username",
            ""
        ).strip()

        password = request.POST.get(
            "password",
            ""
        )

        user = authenticate(
            request,
            username=username,
            password=password
        )

        if user is not None:

            login(request, user)

            return redirect("home")

        return render(
            request,
            "login.html",
            {
                "error":
                    "Invalid username or password."
            }
        )

    return render(request, "login.html")


# ============================================================
# LOGOUT
# ============================================================

def logout_view(request):

    logout(request)

    return redirect("home")


# ============================================================
# HISTORY
# ============================================================

@login_required
def history(request):

    resumes = ResumeAnalysis.objects.filter(
        user=request.user
    ).order_by("-analyzed_at")

    searches = CareerSearch.objects.filter(
        user=request.user
    ).order_by("-searched_at")

    return render(
        request,
        "history.html",
        {
            "resumes": resumes,
            "searches": searches,
        }
    )


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(file_path):

    text = ""

    try:

        reader = PdfReader(file_path)

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    except Exception as e:

        print("PDF EXTRACTION ERROR:", e)

    return text.strip()


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(file_path):

    text = ""

    try:

        document = Document(file_path)

        for paragraph in document.paragraphs:

            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():
                        text += cell.text + "\n"

    except Exception as e:

        print("DOCX EXTRACTION ERROR:", e)

    return text.strip()


# ============================================================
# RESUME EXTRACTION
# ============================================================

def extract_resume_text(file_path):

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    if extension == ".docx":
        return extract_docx_text(file_path)

    return ""


# ============================================================
# NORMALIZE TEXT
# ============================================================

def normalize_text(text):

    text = text.lower()

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


# ============================================================
# SKILL DATABASE
# ============================================================

SKILLS = {

    "Python": ["python"],
    "Java": ["java"],
    "JavaScript": ["javascript", "js"],
    "TypeScript": ["typescript"],
    "HTML": ["html"],
    "CSS": ["css"],
    "React": ["react", "reactjs", "react.js"],
    "Django": ["django"],
    "Flask": ["flask"],
    "FastAPI": ["fastapi"],
    "Node.js": ["node.js", "nodejs"],
    "SQL": ["sql"],
    "MySQL": ["mysql"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MongoDB": ["mongodb", "mongo"],
    "SQLite": ["sqlite"],
    "Git": ["git"],
    "GitHub": ["github"],
    "Docker": ["docker"],
    "AWS": ["aws", "amazon web services"],
    "Azure": ["azure"],
    "Machine Learning": [
        "machine learning",
        "machine-learning"
    ],
    "Deep Learning": [
        "deep learning",
        "deep-learning"
    ],
    "Artificial Intelligence": [
        "artificial intelligence",
        "ai"
    ],
    "Data Science": [
        "data science",
        "data analytics"
    ],
    "Pandas": ["pandas"],
    "NumPy": ["numpy"],
    "TensorFlow": ["tensorflow"],
    "PyTorch": ["pytorch"],
    "Power BI": ["power bi"],
    "Excel": ["excel", "microsoft excel"],
    "C++": ["c++"],
    "C#": ["c#"],
    "Kotlin": ["kotlin"],
    "Swift": ["swift"],
    "PHP": ["php"],
    "REST API": [
        "rest api",
        "restful api"
    ],
    "Linux": ["linux"],
    "Networking": [
        "networking",
        "computer networking"
    ],
    "Cybersecurity": [
        "cybersecurity",
        "cyber security",
        "information security"
    ],
    "Problem Solving": [
        "problem solving",
        "problem-solving"
    ],
    "Communication": ["communication"],
    "Leadership": ["leadership"],
    "Teamwork": [
        "teamwork",
        "team work"
    ]
}


# ============================================================
# CAREER REQUIREMENTS
# ============================================================

CAREER_REQUIREMENTS = {

    "Python Developer": [
        "Python",
        "Django",
        "SQL",
        "Git",
        "REST API"
    ],

    "Web Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "Git"
    ],

    "Full Stack Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "React",
        "Node.js",
        "SQL",
        "Git"
    ],

    "Data Analyst": [
        "Python",
        "SQL",
        "Excel",
        "Power BI",
        "Pandas"
    ],

    "Data Scientist": [
        "Python",
        "SQL",
        "Pandas",
        "NumPy",
        "Machine Learning"
    ],

    "Machine Learning Engineer": [
        "Python",
        "Machine Learning",
        "Deep Learning",
        "TensorFlow",
        "SQL"
    ],

    "Backend Developer": [
        "Python",
        "Django",
        "SQL",
        "REST API",
        "Git"
    ],

    "Frontend Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "React",
        "Git"
    ],

    "Cybersecurity Analyst": [
        "Cybersecurity",
        "Linux",
        "Networking",
        "Python"
    ]
}


# ============================================================
# DETECT SKILLS
# ============================================================

def detect_skills(text):

    normalized = normalize_text(text)

    detected = []

    for skill, keywords in SKILLS.items():

        for keyword in keywords:

            keyword_normalized = normalize_text(keyword)

            pattern = (
                r"(?<!\w)"
                + re.escape(keyword_normalized)
                + r"(?!\w)"
            )

            if re.search(
                pattern,
                normalized
            ):

                detected.append(skill)
                break

    return detected


# ============================================================
# CAREER MATCHING
# ============================================================

def calculate_career_matches(detected_skills):

    detected_set = set(detected_skills)

    results = []

    for career, required_skills in CAREER_REQUIREMENTS.items():

        matched = [
            skill
            for skill in required_skills
            if skill in detected_set
        ]

        missing = [
            skill
            for skill in required_skills
            if skill not in detected_set
        ]

        total = len(required_skills)

        score = (
            round(
                len(matched) / total * 100
            )
            if total
            else 0
        )

        results.append(
            {
                "career": career,
                "score": score,
                "matched": matched,
                "missing": missing
            }
        )

    results.sort(
        key=lambda item: item["score"],
        reverse=True
    )

    return results


# ============================================================
# RESUME SUMMARY
# ============================================================

def create_summary(text, detected_skills):

    word_count = len(text.split())

    if word_count == 0:

        return (
            "No readable text was detected "
            "from this resume."
        )

    if detected_skills:

        skill_text = ", ".join(
            detected_skills[:8]
        )

        return (
            f"Your resume contains approximately "
            f"{word_count} words. The analyzer detected "
            f"skills including {skill_text}. "
            f"Overall, your resume contains "
            f"{len(detected_skills)} recognized skills."
        )

    return (
        f"Your resume contains approximately "
        f"{word_count} words. The analyzer could not "
        f"confidently identify standard technical skills. "
        f"Consider making your skills and technologies "
        f"more explicit."
    )


# ============================================================
# RESUME QUALITY
# ============================================================

def analyze_resume_quality(text):

    normalized = normalize_text(text)

    strengths = []
    suggestions = []

    email_found = bool(
        re.search(
            r"[\w\.-]+@[\w\.-]+\.\w+",
            text
        )
    )

    if email_found:
        strengths.append(
            "Professional email address detected."
        )
    else:
        suggestions.append(
            "Add a professional email address."
        )

    phone_found = bool(
        re.search(
            r"(\+?\d[\d\s\-\(\)]{8,}\d)",
            text
        )
    )

    if phone_found:
        strengths.append(
            "Phone number detected."
        )
    else:
        suggestions.append(
            "Add a phone number."
        )

    education_words = [
        "education",
        "university",
        "college",
        "degree",
        "bachelor",
        "master"
    ]

    if any(
        word in normalized
        for word in education_words
    ):

        strengths.append(
            "Education information detected."
        )

    else:

        suggestions.append(
            "Add a clear Education section."
        )

    experience_words = [
        "experience",
        "employment",
        "work history",
        "internship"
    ]

    if any(
        word in normalized
        for word in experience_words
    ):

        strengths.append(
            "Experience information detected."
        )

    else:

        suggestions.append(
            "Add work experience or internship details."
        )

    if "project" in normalized:

        strengths.append(
            "Project information detected."
        )

    else:

        suggestions.append(
            "Add 2–3 relevant projects with technologies used."
        )

    if (
        "skills" in normalized
        or "technical skills" in normalized
    ):

        strengths.append(
            "Skills section detected."
        )

    else:

        suggestions.append(
            "Create a dedicated Technical Skills section."
        )

    action_verbs = [
        "developed",
        "created",
        "built",
        "implemented",
        "designed",
        "managed",
        "led",
        "improved",
        "optimized",
        "automated",
        "analyzed",
        "deployed"
    ]

    action_count = 0

    for verb in action_verbs:

        action_count += len(
            re.findall(
                r"\b" + verb + r"\b",
                normalized
            )
        )

    if action_count >= 3:

        strengths.append(
            "Good use of action-oriented language."
        )

    else:

        suggestions.append(
            "Use stronger action verbs such as "
            "developed, implemented, designed, "
            "optimized, and automated."
        )

    numbers = re.findall(
        r"\b\d+(?:\.\d+)?%?\b",
        text
    )

    if len(numbers) >= 5:

        strengths.append(
            "Resume contains measurable details."
        )

    else:

        suggestions.append(
            "Add measurable achievements such as "
            "percentages, users, time saved, "
            "performance improvements, or results."
        )

    word_count = len(text.split())

    if word_count < 150:

        suggestions.append(
            "Your resume appears short. Add meaningful "
            "project, education, and experience details."
        )

    elif word_count > 1500:

        suggestions.append(
            "Your resume may be longer than necessary. "
            "Remove repetitive information."
        )

    else:

        strengths.append(
            "Resume length appears reasonable."
        )

    return strengths, suggestions


# ============================================================
# INTERVIEW QUESTIONS
# ============================================================

def create_interview_questions(
    detected_skills,
    career
):

    questions = [

        (
            f"Tell me about yourself and why you are "
            f"interested in a {career} role."
        ),

        "What is your strongest technical skill?",

        "Tell me about your most important project."
    ]

    for skill in detected_skills[:5]:

        questions.append(
            f"Explain how you have used {skill} "
            f"in a project or practical situation."
        )

    questions.extend(
        [
            (
                "Describe a difficult technical problem "
                "you faced and how you solved it."
            ),
            (
                "What technical skill are you currently "
                "working to improve?"
            )
        ]
    )

    return questions


# ============================================================
# GEMMA RESUME ANALYSIS
# ============================================================

def generate_resume_ai_analysis(
    resume_text,
    detected_skills,
    career_matches,
    top_career,
    missing_skills,
    strengths
):

    match_text = ""

    for match in career_matches[:5]:

        match_text += (
            f"{match['career']}: "
            f"{match['score']}%, "
            f"matched={match['matched']}, "
            f"missing={match['missing']}\n"
        )

    prompt = f"""
You are CareerAI.

Analyze this student's resume and recommend the
most suitable entry-level technology career.

RESUME:
{resume_text[:12000]}

DETECTED SKILLS:
{", ".join(detected_skills)}

RULE-BASED CAREER MATCHES:
{match_text}

RULE-BASED TOP CAREER:
{top_career}

MISSING SKILLS:
{", ".join(missing_skills)}

RESUME STRENGTHS:
{", ".join(strengths)}

Return ONLY valid JSON.

Use exactly this structure:

{{
    "career": "recommended career",
    "match_percentage": 0,
    "why": "Short explanation of why this career fits the resume.",
    "good_fit": [
        "reason 1",
        "reason 2",
        "reason 3"
    ],
    "skills_to_improve": [
        "skill 1",
        "skill 2",
        "skill 3"
    ],
    "learning_path": [
        "step 1",
        "step 2",
        "step 3",
        "step 4",
        "step 5"
    ],
    "resume_summary": "A professional short resume summary.",
    "interview_questions": [
        "question 1",
        "question 2",
        "question 3",
        "question 4",
        "question 5"
    ]
}}

Important:

- Use only information supported by the resume.
- Do not invent work experience.
- Do not claim the student has skills that are not shown.
- The match percentage should be realistic.
- Do not include markdown.
- Do not include code fences.
"""

    response = ask_gemma(prompt)

    if not response:
        return None

    try:

        cleaned = clean_gemma_json(response)

        return json.loads(cleaned)

    except Exception as e:

        print("GEMMA RESUME JSON ERROR:", e)
        print("GEMMA RESPONSE:", response)

        return None


# ============================================================
# RESUME ANALYZER
# ============================================================

@login_required
def resume_analyzer(request):

    if request.method != "POST":

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": None
            }
        )

    uploaded_file = request.FILES.get("resume")

    if not uploaded_file:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": "Please select a resume file."
            }
        )

    extension = Path(
        uploaded_file.name
    ).suffix.lower()

    allowed_extensions = [
        ".pdf",
        ".docx"
    ]

    if extension not in allowed_extensions:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error":
                    "Please upload a PDF or DOCX file."
            }
        )

    # --------------------------------------------------------
    # SAVE RESUME
    # --------------------------------------------------------

    resume_analysis = ResumeAnalysis.objects.create(
        user=request.user,
        resume=uploaded_file
    )

    # --------------------------------------------------------
    # EXTRACT TEXT
    # --------------------------------------------------------

    try:

        resume_text = extract_resume_text(
            resume_analysis.resume.path
        )

    except Exception as e:

        print("RESUME EXTRACTION ERROR:", e)

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error":
                    "The resume was uploaded, "
                    "but text extraction failed."
            }
        )

    if not resume_text:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error":
                    "The resume was uploaded, "
                    "but no readable text was found."
            }
        )

    # --------------------------------------------------------
    # SAVE TEXT
    # --------------------------------------------------------

    resume_analysis.extracted_text = resume_text
    resume_analysis.save()

    # --------------------------------------------------------
    # LOCAL ANALYSIS
    # --------------------------------------------------------

    detected_skills = detect_skills(
        resume_text
    )

    career_matches = calculate_career_matches(
        detected_skills
    )

    if career_matches:

        top_career = career_matches[0]["career"]
        missing_skills = career_matches[0]["missing"]

    else:

        top_career = "Software Developer"
        missing_skills = []

    strengths, suggestions = analyze_resume_quality(
        resume_text
    )

    interview_questions = create_interview_questions(
        detected_skills,
        top_career
    )

    summary = create_summary(
        resume_text,
        detected_skills
    )

    # --------------------------------------------------------
    # GEMMA ANALYSIS
    # --------------------------------------------------------

    ai_analysis = generate_resume_ai_analysis(
        resume_text=resume_text,
        detected_skills=detected_skills,
        career_matches=career_matches,
        top_career=top_career,
        missing_skills=missing_skills,
        strengths=strengths
    )

    # --------------------------------------------------------
    # AI RESULT
    # --------------------------------------------------------

    if ai_analysis:

        ai_career = ai_analysis.get(
            "career",
            top_career
        )

        ai_match = ai_analysis.get(
            "match_percentage",
            career_matches[0]["score"]
            if career_matches
            else 0
        )

        ai_why = ai_analysis.get(
            "why",
            ""
        )

        ai_good_fit = ai_analysis.get(
            "good_fit",
            strengths
        )

        ai_skills = ai_analysis.get(
            "skills_to_improve",
            missing_skills
        )

        ai_learning_path = ai_analysis.get(
            "learning_path",
            []
        )

        ai_resume_summary = ai_analysis.get(
            "resume_summary",
            summary
        )

        ai_interview_questions = ai_analysis.get(
            "interview_questions",
            interview_questions
        )

    else:

        ai_career = top_career

        ai_match = (
            career_matches[0]["score"]
            if career_matches
            else 0
        )

        ai_why = (
            "The local AI service was unavailable, "
            "so the rule-based career analyzer "
            "was used."
        )

        ai_good_fit = strengths
        ai_skills = missing_skills
        ai_learning_path = []
        ai_resume_summary = summary
        ai_interview_questions = interview_questions

    # --------------------------------------------------------
    # FINAL ANALYSIS
    # --------------------------------------------------------

    analysis = {

        "resume": resume_analysis,

        "summary": summary,

        "skills": detected_skills,

        "missing_skills": missing_skills,

        "career_matches": career_matches[:5],

        "top_career": ai_career,

        "strengths": strengths,

        "suggestions": suggestions,

        "interview_questions":
            ai_interview_questions,

        "word_count":
            len(resume_text.split()),

        "ai_available":
            bool(ai_analysis),

        "ai_powered":
            bool(ai_analysis),

        "career_score":
            ai_match,

        "career_reason":
            ai_why,

        "career_strengths":
            ai_good_fit,

        "learning_path":
            ai_learning_path,

        "ai_career":
            ai_career,

        "ai_match":
            ai_match,

        "ai_why":
            ai_why,

        "ai_good_fit":
            ai_good_fit,

        "ai_skills":
            ai_skills,

        "ai_learning_path":
            ai_learning_path,

        "ai_resume_summary":
            ai_resume_summary,

        "ai_interview_questions":
            ai_interview_questions,
    }

    return render(
        request,
        "resume.html",
        {
            "analysis": analysis,
            "error": None
        }
    )