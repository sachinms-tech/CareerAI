from pathlib import Path
import re

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.shortcuts import redirect, render

from pypdf import PdfReader
from docx import Document

from .models import CareerSearch, ResumeAnalysis


# ============================================================
# HOME
# ============================================================

def home(request):
    return render(request, "home.html")


# ============================================================
# REGISTER
# ============================================================

def register_view(request):

    if request.method == "POST":

        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip()
        password = request.POST.get("password", "")
        confirm_password = request.POST.get("confirm_password", "")

        if not username or not password:
            return render(
                request,
                "register.html",
                {"error": "Username and password are required."}
            )

        if password != confirm_password:
            return render(
                request,
                "register.html",
                {"error": "Passwords do not match."}
            )

        if User.objects.filter(username=username).exists():
            return render(
                request,
                "register.html",
                {"error": "Username already exists."}
            )

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password
        )

        login(request, user)

        return redirect("home")

    return render(request, "register.html")


# ============================================================
# LOGIN
# ============================================================

def login_view(request):

    if request.method == "POST":

        username = request.POST.get("username", "").strip()
        password = request.POST.get("password", "")

        user = authenticate(
            request,
            username=username,
            password=password
        )

        if user is not None:
            login(request, user)
            return redirect("home")

        return render(
            request,
            "login.html",
            {"error": "Invalid username or password."}
        )

    return render(request, "login.html")


# ============================================================
# LOGOUT
# ============================================================

def logout_view(request):

    logout(request)

    return redirect("home")


# ============================================================
# HISTORY
# ============================================================

@login_required
def history(request):

    resumes = ResumeAnalysis.objects.filter(
        user=request.user
    ).order_by("-analyzed_at")

    searches = CareerSearch.objects.filter(
        user=request.user
    ).order_by("-searched_at")

    return render(
        request,
        "history.html",
        {
            "resumes": resumes,
            "searches": searches,
        }
    )


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(file_path):

    text = ""

    try:
        reader = PdfReader(file_path)

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    except Exception as e:
        print("PDF EXTRACTION ERROR:", e)

    return text.strip()


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(file_path):

    text = ""

    try:
        document = Document(file_path)

        for paragraph in document.paragraphs:

            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():
                        text += cell.text + "\n"

    except Exception as e:
        print("DOCX EXTRACTION ERROR:", e)

    return text.strip()


# ============================================================
# RESUME EXTRACTION
# ============================================================

def extract_resume_text(file_path):

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    if extension == ".docx":
        return extract_docx_text(file_path)

    return ""


# ============================================================
# NORMALIZE TEXT
# ============================================================

def normalize_text(text):

    text = text.lower()

    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ============================================================
# SKILL DATABASE
# ============================================================

SKILLS = {

    "Python": ["python"],
    "Java": ["java"],
    "JavaScript": ["javascript", "js"],
    "TypeScript": ["typescript"],
    "HTML": ["html"],
    "CSS": ["css"],

    "React": [
        "react",
        "reactjs",
        "react.js"
    ],

    "Django": ["django"],
    "Flask": ["flask"],
    "FastAPI": ["fastapi"],

    "Node.js": [
        "node.js",
        "nodejs"
    ],

    "SQL": ["sql"],
    "MySQL": ["mysql"],

    "PostgreSQL": [
        "postgresql",
        "postgres"
    ],

    "MongoDB": [
        "mongodb",
        "mongo"
    ],

    "SQLite": ["sqlite"],

    "Git": ["git"],
    "GitHub": ["github"],
    "Docker": ["docker"],

    "AWS": [
        "aws",
        "amazon web services"
    ],

    "Azure": ["azure"],

    "Machine Learning": [
        "machine learning",
        "machine-learning"
    ],

    "Deep Learning": [
        "deep learning",
        "deep-learning"
    ],

    "Artificial Intelligence": [
        "artificial intelligence",
        "ai"
    ],

    "Data Science": [
        "data science",
        "data analytics"
    ],

    "Pandas": ["pandas"],
    "NumPy": ["numpy"],
    "TensorFlow": ["tensorflow"],
    "PyTorch": ["pytorch"],

    "Power BI": ["power bi"],

    "Excel": [
        "excel",
        "microsoft excel"
    ],

    "C++": ["c++"],
    "C#": ["c#"],
    "Kotlin": ["kotlin"],
    "Swift": ["swift"],
    "PHP": ["php"],

    "REST API": [
        "rest api",
        "restful api"
    ],

    "Linux": ["linux"],

    "Networking": [
        "networking",
        "computer networking"
    ],

    "Cybersecurity": [
        "cybersecurity",
        "cyber security",
        "information security"
    ],

    "Problem Solving": [
        "problem solving",
        "problem-solving"
    ],

    "Communication": [
        "communication"
    ],

    "Leadership": [
        "leadership"
    ],

    "Teamwork": [
        "teamwork",
        "team work"
    ]
}


# ============================================================
# CAREER REQUIREMENTS
# ============================================================

CAREER_REQUIREMENTS = {

    "Python Developer": [
        "Python",
        "Django",
        "SQL",
        "Git",
        "REST API"
    ],

    "Web Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "Git"
    ],

    "Full Stack Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "React",
        "Node.js",
        "SQL",
        "Git"
    ],

    "Data Analyst": [
        "Python",
        "SQL",
        "Excel",
        "Power BI",
        "Pandas"
    ],

    "Data Scientist": [
        "Python",
        "SQL",
        "Pandas",
        "NumPy",
        "Machine Learning"
    ],

    "Machine Learning Engineer": [
        "Python",
        "Machine Learning",
        "Deep Learning",
        "TensorFlow",
        "SQL"
    ],

    "Backend Developer": [
        "Python",
        "Django",
        "SQL",
        "REST API",
        "Git"
    ],

    "Frontend Developer": [
        "HTML",
        "CSS",
        "JavaScript",
        "React",
        "Git"
    ],

    "Cybersecurity Analyst": [
        "Cybersecurity",
        "Linux",
        "Networking",
        "Python"
    ]
}


# ============================================================
# DETECT SKILLS
# ============================================================

def detect_skills(text):

    normalized = normalize_text(text)

    detected = []

    for skill, keywords in SKILLS.items():

        for keyword in keywords:

            keyword_normalized = normalize_text(keyword)

            pattern = (
                r"(?<!\w)"
                + re.escape(keyword_normalized)
                + r"(?!\w)"
            )

            if re.search(pattern, normalized):

                detected.append(skill)

                break

    return detected


# ============================================================
# CAREER MATCHING
# ============================================================

def calculate_career_matches(detected_skills):

    detected_set = set(detected_skills)

    results = []

    for career, required_skills in CAREER_REQUIREMENTS.items():

        matched = [
            skill
            for skill in required_skills
            if skill in detected_set
        ]

        missing = [
            skill
            for skill in required_skills
            if skill not in detected_set
        ]

        total = len(required_skills)

        score = round(
            (len(matched) / total) * 100
        ) if total else 0

        results.append({
            "career": career,
            "score": score,
            "matched": matched,
            "missing": missing
        })

    results.sort(
        key=lambda item: item["score"],
        reverse=True
    )

    return results


# ============================================================
# SUMMARY
# ============================================================

def create_summary(text, detected_skills):

    word_count = len(text.split())

    if word_count == 0:

        return (
            "No readable text was detected "
            "from this resume."
        )

    if detected_skills:

        skill_text = ", ".join(
            detected_skills[:8]
        )

        return (
            f"Your resume contains approximately "
            f"{word_count} words. The analyzer detected "
            f"skills including {skill_text}. "
            f"Overall, your resume contains "
            f"{len(detected_skills)} recognized skills."
        )

    return (
        f"Your resume contains approximately "
        f"{word_count} words. The analyzer could not "
        f"confidently identify standard technical skills. "
        f"Consider making your skills and technologies "
        f"more explicit."
    )


# ============================================================
# RESUME QUALITY
# ============================================================

def analyze_resume_quality(text):

    normalized = normalize_text(text)

    strengths = []
    suggestions = []

    # EMAIL
    email_found = bool(
        re.search(
            r"[\w\.-]+@[\w\.-]+\.\w+",
            text
        )
    )

    if email_found:
        strengths.append(
            "Professional email address detected."
        )
    else:
        suggestions.append(
            "Add a professional email address."
        )

    # PHONE
    phone_found = bool(
        re.search(
            r"(\+?\d[\d\s\-\(\)]{8,}\d)",
            text
        )
    )

    if phone_found:
        strengths.append(
            "Phone number detected."
        )
    else:
        suggestions.append(
            "Add a phone number."
        )

    # EDUCATION
    education_words = [
        "education",
        "university",
        "college",
        "degree",
        "bachelor",
        "master"
    ]

    if any(
        word in normalized
        for word in education_words
    ):
        strengths.append(
            "Education information detected."
        )
    else:
        suggestions.append(
            "Add a clear Education section."
        )

    # EXPERIENCE
    experience_words = [
        "experience",
        "employment",
        "work history",
        "internship"
    ]

    if any(
        word in normalized
        for word in experience_words
    ):
        strengths.append(
            "Experience information detected."
        )
    else:
        suggestions.append(
            "Add work experience or internship details."
        )

    # PROJECTS
    if "project" in normalized:
        strengths.append(
            "Project information detected."
        )
    else:
        suggestions.append(
            "Add 2–3 relevant projects with technologies used."
        )

    # SKILLS
    if (
        "skills" in normalized
        or "technical skills" in normalized
    ):
        strengths.append(
            "Skills section detected."
        )
    else:
        suggestions.append(
            "Create a dedicated Technical Skills section."
        )

    # ACTION VERBS
    action_verbs = [
        "developed",
        "created",
        "built",
        "implemented",
        "designed",
        "managed",
        "led",
        "improved",
        "optimized",
        "automated",
        "analyzed",
        "deployed"
    ]

    action_count = 0

    for verb in action_verbs:

        action_count += len(
            re.findall(
                r"\b" + verb + r"\b",
                normalized
            )
        )

    if action_count >= 3:
        strengths.append(
            "Good use of action-oriented language."
        )
    else:
        suggestions.append(
            "Use stronger action verbs such as "
            "developed, implemented, designed, "
            "optimized, and automated."
        )

    # MEASURABLE ACHIEVEMENTS
    numbers = re.findall(
        r"\b\d+(?:\.\d+)?%?\b",
        text
    )

    if len(numbers) >= 5:
        strengths.append(
            "Resume contains measurable details."
        )
    else:
        suggestions.append(
            "Add measurable achievements such as "
            "percentages, users, time saved, "
            "performance improvements, or results."
        )

    # WORD COUNT
    word_count = len(text.split())

    if word_count < 150:

        suggestions.append(
            "Your resume appears short. Add meaningful "
            "project, education, and experience details."
        )

    elif word_count > 1500:

        suggestions.append(
            "Your resume may be longer than necessary. "
            "Remove repetitive information."
        )

    else:

        strengths.append(
            "Resume length appears reasonable."
        )

    return strengths, suggestions


# ============================================================
# INTERVIEW QUESTIONS
# ============================================================

def create_interview_questions(
    detected_skills,
    career
):

    questions = [
        (
            f"Tell me about yourself and why you are "
            f"interested in a {career} role."
        ),
        "What is your strongest technical skill?",
        "Tell me about your most important project.",
    ]

    for skill in detected_skills[:5]:

        questions.append(
            f"Explain how you have used {skill} "
            f"in a project or practical situation."
        )

    questions.extend([
        (
            "Describe a difficult technical problem "
            "you faced and how you solved it."
        ),
        (
            "What technical skill are you currently "
            "working to improve?"
        )
    ])

    return questions


# ============================================================
# RESUME ANALYZER
# ============================================================

@login_required
def resume_analyzer(request):

    if request.method != "POST":

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": None
            }
        )

    uploaded_file = request.FILES.get("resume")

    if not uploaded_file:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": "Please select a resume file."
            }
        )

    extension = Path(
        uploaded_file.name
    ).suffix.lower()

    allowed_extensions = [
        ".pdf",
        ".docx"
    ]

    if extension not in allowed_extensions:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": (
                    "Please upload a PDF or DOCX file."
                )
            }
        )

    # SAVE
    resume_analysis = ResumeAnalysis.objects.create(
        user=request.user,
        resume=uploaded_file
    )

    # EXTRACT
    try:

        resume_text = extract_resume_text(
            resume_analysis.resume.path
        )

    except Exception as e:

        print(
            "RESUME EXTRACTION ERROR:",
            e
        )

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": (
                    "The resume was uploaded, "
                    "but text extraction failed."
                )
            }
        )

    if not resume_text:

        return render(
            request,
            "resume.html",
            {
                "analysis": None,
                "error": (
                    "The resume was uploaded, "
                    "but no readable text was found."
                )
            }
        )

    # SAVE TEXT
    resume_analysis.extracted_text = resume_text
    resume_analysis.save()

    # LOCAL ANALYSIS
    detected_skills = detect_skills(
        resume_text
    )

    career_matches = calculate_career_matches(
        detected_skills
    )

    if career_matches:

        top_career = career_matches[0]["career"]

        missing_skills = (
            career_matches[0]["missing"]
        )

    else:

        top_career = "Software Developer"

        missing_skills = []

    strengths, suggestions = (
        analyze_resume_quality(
            resume_text
        )
    )

    interview_questions = (
        create_interview_questions(
            detected_skills,
            top_career
        )
    )

    summary = create_summary(
        resume_text,
        detected_skills
    )

    analysis = {

        "resume": resume_analysis,

        "summary": summary,

        "skills": detected_skills,

        "missing_skills": missing_skills,

        "career_matches": career_matches[:5],

        "top_career": top_career,

        "strengths": strengths,

        "suggestions": suggestions,

        "interview_questions": interview_questions,

        "word_count": len(
            resume_text.split()
        )
    }

    return render(
        request,
        "resume.html",
        {
            "analysis": analysis,
            "error": None
        }
    )